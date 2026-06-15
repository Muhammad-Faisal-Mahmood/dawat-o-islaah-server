
# hadith/admin.py
# FULL UPDATED FILE
# Copy paste complete file

import re

from django.contrib import admin
from django import forms
from ckeditor.widgets import CKEditorWidget
from .models import HadithDocumentUpload, Hadith, Baab, Book, Chapter
from django.urls import path
from django.http import HttpResponse
from django.db import transaction




# ======================================================
# FORM
# ======================================================
class HadithAdminForm(forms.ModelForm):
    arabic_text = forms.CharField(widget=CKEditorWidget(), required=False)
    english_text = forms.CharField(widget=CKEditorWidget(), required=False)
    reference = forms.CharField(widget=CKEditorWidget(), required=False)
    urdu_text = forms.CharField(widget=CKEditorWidget(), required=False)

    class Meta:
        model = Hadith
        fields = "__all__"


# ======================================================
# BOOK ADMIN
# ======================================================
@admin.register(Book)
class BookAdmin(admin.ModelAdmin):
    list_display = ("name", "order")
    ordering = ("order",)


# ======================================================
# CHAPTER ADMIN
# ======================================================
@admin.register(Chapter)
class ChapterAdmin(admin.ModelAdmin):
    list_display = ("chapter_number", "chapter_english", "book")
    list_filter = ("book",)
    search_fields = ("chapter_english", "chapter_urdu", "chapter_arabic")
    ordering = ("book", "chapter_number")


# ======================================================
# BAAB ADMIN FORM
# ======================================================
class BaabAdminForm(forms.ModelForm):
    chapter = forms.ModelChoiceField(
        queryset=Chapter.objects.all(),
        required=False,
        widget=forms.Select(attrs={"class": "chapter-select"})
    )

    class Meta:
        model = Baab
        fields = "__all__"

    class Media:
        js = ("admin/js/baab_admin.js",)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        book_id = None
        if self.instance.pk:
            book_id = self.instance.book_id
        elif self.is_bound:
            try:
                book_id = int(self.data.get("book"))
            except (ValueError, TypeError):
                pass

        if book_id:
            self.fields["chapter"].queryset = Chapter.objects.filter(
                book=book_id
            ).order_by("chapter_number")
        else:
            self.fields["chapter"].queryset = Chapter.objects.none()

    URDU_RE = re.compile(
        r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]'
    )

    def clean_baab_name_urdu(self):
        value = self.cleaned_data.get("baab_name_urdu", "")
        if value and not self.URDU_RE.search(value):
            raise forms.ValidationError(
                "This field appears to be in English/Latin script. "
                "Please enter the baab name in Urdu script."
            )
        return value

    def clean_baab_name_english(self):
        value = self.cleaned_data.get("baab_name_english", "")
        if value and self.URDU_RE.search(value):
            raise forms.ValidationError(
                "This field appears to be in Urdu/Arabic script. "
                "Please enter the baab name in English/Latin script."
            )
        return value


# ======================================================
# BAAB ADMIN
# ======================================================
@admin.register(Baab)
class BaabAdmin(admin.ModelAdmin):
    form = BaabAdminForm

    list_display = (
        "baab_name_urdu",
        "baab_name_english",
        "book",
        "get_chapter_display",
        "start_hadith_number",
        "end_hadith_number",
    )

    list_filter = ("book",)

    fieldsets = (
        ("Book & Chapter", {
            "fields": ("book", "chapter")
        }),
        ("Baab Names", {
            "fields": ("baab_name_urdu", "baab_name_english")
        }),
        ("Hadith Range", {
            "fields": ("start_hadith_number", "end_hadith_number")
        }),
    )

    @admin.display(description="Chapter")
    def get_chapter_display(self, obj):
        if obj.chapter:
            return f"{obj.chapter.chapter_number}. {obj.chapter.chapter_english}"
        return obj.chapter_english or "-"

    def save_model(self, request, obj, form, change):
        from hadith.models import Hadith
        super().save_model(request, obj, form, change)
        updated = Hadith.objects.filter(
            book=obj.book,
            hadith_number__gte=obj.start_hadith_number,
            hadith_number__lte=obj.end_hadith_number
        ).exclude(baab=obj).update(baab=obj)
        if updated:
            self.message_user(request, f"Linked {updated} hadith(s) to this baab.")


# ======================================================
# CHAPTER FILTER
# ======================================================
class ChapterNumberFilter(admin.SimpleListFilter):
    title = "Chapter"
    parameter_name = "chapter_num"

    def lookups(self, request, model_admin):
        qs = Hadith.objects.all()

        book_id = request.GET.get("book__id__exact")
        if book_id:
            qs = qs.filter(book_id=book_id)

        chapters = (
            qs.order_by("chapter_number")
            .values_list("chapter_number", "chapter_english")
            .distinct()
        )

        seen = set()
        results = []

        for num, name in chapters:
            if num not in seen:
                seen.add(num)
                results.append((str(num), f"{num}. {name}"))

        return results

    def queryset(self, request, queryset):
        if self.value():
            return queryset.filter(chapter_number=int(self.value()))
        return queryset


# ======================================================
# HADITH ADMIN
# ======================================================
@admin.register(Hadith)
class HadithAdmin(admin.ModelAdmin):
    form = HadithAdminForm

    # GLOBAL HADITH NUMBER DISPLAY
    list_display = (
        "chapter_hadith_id",
        "book",
        "chapter_number",
        "chapter_english",
        "chapter_urdu",
        "chapter_arabic",
        "get_baab",
        "status",
        "detailed_explanation",
    )

    list_display_links = ("chapter_hadith_id",)

    search_fields = (
        "chapter_hadith_id",
        "hadith_number",
        "chapter_english",
        "chapter_urdu",
        "chapter_arabic",
        "reference",
        "status",
        "detailed_explanation",
        "baab__baab_name_urdu",
        "baab__baab_name_english",
        "book__name",
    )

    list_filter = (
        "book",
        ChapterNumberFilter,
    )

    # IMPORTANT FIX
    # use global hadith_number instead of chapter_hadith_id
    ordering = (
        "book__order",
        "chapter_hadith_id",
    )

    list_per_page = 50

    fieldsets = (
        ("Main Info", {
            "fields": (
                "book",
                "chapter_number",
                "chapter_hadith_id",
                "hadith_number",
                "baab",
            )
        }),

        ("Chapter Names", {
            "fields": (
                "chapter_english",
                "chapter_urdu",
                "chapter_arabic",
            )
        }),

        ("Texts", {
            "fields": (
                "arabic_text",
                "urdu_text",
                "english_text",
            )
        }),

        ("Other", {
            "fields": (
                "reference",
                "updated_by",
                "status",
                "detailed_explanation",
            )
        }),
    )

    def get_queryset(self, request):
        return super().get_queryset(request).select_related(
            "book",
            "baab"
        )

    @admin.display(description="Baab Urdu")
    def get_baab(self, obj):
        return obj.baab.baab_name_urdu if obj.baab else "-"

    @admin.display(description="Baab English")
    def get_baab_english(self, obj):
        return obj.baab.baab_name_english if obj.baab and obj.baab.baab_name_english else "-"

@admin.register(HadithDocumentUpload)
class HadithDocumentUploadAdmin(admin.ModelAdmin):
    readonly_fields = ('result_log', 'uploaded_at')
    list_display = ('__str__', 'uploaded_at', 'result_summary')
    change_list_template = 'admin/hadith/hadithdocumentupload/change_list.html'

    fieldsets = (
        (None, {
            'fields': ('file',)
        }),
        ('Results', {
            'fields': ('result_log', 'uploaded_at'),
            'classes': ('wide',),
        }),
    )

    def result_summary(self, obj):
        if obj.result_log:
            if '???' in obj.result_log:
                return obj.result_log[:120]
            lines = [l for l in obj.result_log.split('\n') if l.strip()]
            return lines[0] if lines else '...'
        return 'Pending...'
    result_summary.short_description = 'Result'

    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('download-template/', self.admin_site.admin_view(self.download_template),
                 name='hadith_hadithdocumentupload_download_template'),
        ]
        return custom_urls + urls

    def download_template(self, request):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from io import BytesIO

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        doc.add_heading('Chapter Information', level=2)
        tbl1 = doc.add_table(rows=6, cols=2, style='Table Grid')
        tbl1.alignment = WD_TABLE_ALIGNMENT.LEFT
        h = tbl1.rows[0]
        h.cells[0].text = 'Field'
        h.cells[1].text = 'Value'
        for cell in h.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        fields = ['Book Name', 'Chapter Number', 'Chapter English', 'Chapter Urdu', 'Chapter Arabic']
        for i, f in enumerate(fields):
            tbl1.rows[i+1].cells[0].text = f
        for row in tbl1.rows:
            row.cells[0].width = Cm(4.5)
            row.cells[1].width = Cm(12)

        doc.add_paragraph()
        doc.add_heading('Hadiths', level=2)
        headers = ['Hadith Number', 'Hadith Arabic', 'Hadith Urdu', 'Hadith English',
                   'Reference', 'Status', 'Detailed Explanation']
        tbl2 = doc.add_table(rows=2, cols=7, style='Table Grid')
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        for j, h_text in enumerate(headers):
            cell = tbl2.rows[0].cells[j]
            cell.text = h_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
        for j in range(7):
            tbl2.rows[1].cells[j].text = ''
            for p in tbl2.rows[1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
        widths = [Cm(2), Cm(5.5), Cm(4.5), Cm(5.5), Cm(3), Cm(1.8), Cm(5)]
        for row in tbl2.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="hadith_template.docx"'
        return response

    def save_model(self, request, obj, form, change):
        is_new = not obj.pk
        if is_new:
            parsed = self._parse_upload_document(obj.file)
            if parsed is None:
                self.message_user(
                    request,
                    "This document has more number of hadiths than the number "
                    "of hadiths available for the mentioned chapter in the document. "
                    "Please update the document and match the exact number "
                    "of hadiths the chapter has.",
                    level='ERROR',
                )
                return
            if isinstance(parsed, str):
                self.message_user(request, f"Error: {parsed}", level='ERROR')
                return
        super().save_model(request, obj, form, change)
        if is_new:
            self.process_document(request, obj, parsed)

    def _parse_upload_document(self, file_obj):
        """Parse a .docx upload file.

        Returns a dict of parsed data on success.
        Returns None if the count validation fails.
        Returns an error string for other errors.
        """
        from docx import Document as DocxDocument
        from .models import Book, Chapter, Hadith
        from io import BytesIO

        try:
            buf = BytesIO(file_obj.read())
            file_obj.seek(0)
            doc = DocxDocument(buf)
            tables = doc.tables

            if len(tables) < 2:
                return "Document must contain two tables: 'Chapter Information' and 'Hadiths'"

            # ---- Parse Table 1: Chapter Information ----
            chapter_info = {}
            for row in tables[0].rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0]:
                    chapter_info[cells[0]] = cells[1]

            book_name = chapter_info.get('Book Name', '').strip()
            if not book_name:
                return "'Book Name' is required in Chapter Information table"

            try:
                chapter_number = int(chapter_info.get('Chapter Number', '').strip())
            except (ValueError, TypeError):
                return "'Chapter Number' must be a valid integer"

            # Look up Book - try multiple name formats
            def _resolve_upload_book(name):
                if not name:
                    return None
                n = name.strip()
                book = Book.objects.filter(name__iexact=n).first()
                if book:
                    return book
                slug = n.lower().replace(' ', '-')
                book = Book.objects.filter(name__iexact=slug).first()
                if book:
                    return book
                with_spaces = n.lower().replace('-', ' ')
                book = Book.objects.filter(name__iexact=with_spaces).first()
                if book:
                    return book
                title = n.replace('-', ' ').title()
                book = Book.objects.filter(name__iexact=title).first()
                if book:
                    return book
                book = Book.objects.filter(name__icontains=with_spaces).first()
                if book:
                    return book
                return None

            book = _resolve_upload_book(book_name)
            if not book:
                available = ', '.join(Book.objects.values_list('name', flat=True))
                return f"Book '{book_name}' not found. Available books: {available}"

            chapter_english = chapter_info.get('Chapter English', '').strip()
            chapter_urdu = chapter_info.get('Chapter Urdu', '').strip()
            chapter_arabic = chapter_info.get('Chapter Arabic', '').strip()

            # ---- Parse Table 2: Hadiths ----
            hadith_table = tables[1]
            raw_headers = [
                cell.text.strip() for cell in hadith_table.rows[0].cells
            ]
            rows = []
            for row in hadith_table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 1 and cells[0]:
                    row_data = {}
                    for i, h in enumerate(raw_headers):
                        row_data[h] = cells[i] if i < len(cells) else ''
                    rows.append(row_data)

            if not rows:
                return "No hadith data found in the Hadiths table"

            col_map = {}
            for h in raw_headers:
                hl = h.lower().strip()
                if hl in ('hadith number', 'hadith no', 'no', 'number'):
                    col_map['number'] = h
                elif hl in ('hadith arabic', 'arabic', 'arabic text'):
                    col_map['arabic'] = h
                elif hl in ('hadith urdu', 'urdu', 'urdu text'):
                    col_map['urdu'] = h
                elif hl in ('hadith english', 'english', 'english text'):
                    col_map['english'] = h
                elif hl == 'reference':
                    col_map['reference'] = h
                elif hl == 'status':
                    col_map['status'] = h
                elif hl in ('detailed explanation', 'explanation', 'detail'):
                    col_map['explanation'] = h

            if 'number' not in col_map:
                return (
                    "Could not find 'Hadith Number' column in the Hadiths table. "
                    "Columns found: " + ', '.join(raw_headers)
                )

            # ---- Validate hadith count ----
            existing_count = Hadith.objects.filter(
                book=book, chapter_number=chapter_number
            ).count()
            if len(rows) > existing_count:
                return None  # count mismatch ??? caller will show popup

            return {
                'book': book,
                'chapter_number': chapter_number,
                'chapter_english': chapter_english,
                'chapter_urdu': chapter_urdu,
                'chapter_arabic': chapter_arabic,
                'rows': rows,
                'col_map': col_map,
                'raw_headers': raw_headers,
            }

        except Exception as e:
            return str(e)


    def process_document(self, request, obj, parsed):
        from .models import Book, Chapter, Hadith
        ModelClass = type(obj)

        try:
            book = parsed['book']
            chapter_number = parsed['chapter_number']
            chapter_english = parsed['chapter_english']
            chapter_urdu = parsed['chapter_urdu']
            chapter_arabic = parsed['chapter_arabic']
            rows = parsed['rows']
            col_map = parsed['col_map']

            Chapter.objects.get_or_create(
                book=book,
                chapter_number=chapter_number,
                defaults={
                    'chapter_english': chapter_english or '',
                    'chapter_urdu': chapter_urdu or '',
                    'chapter_arabic': chapter_arabic or '',
                }
            )

            updated_count = 0
            created_count = 0

            with transaction.atomic():
                for row_data in rows:
                    try:
                        hadith_no = int(row_data.get(col_map['number'], '').strip())
                    except (ValueError, TypeError):
                        continue

                    defaults = {
                        'chapter_hadith_id': hadith_no,
                        'chapter_number': chapter_number,
                        'chapter_english': chapter_english or '',
                        'chapter_urdu': chapter_urdu or '',
                        'chapter_arabic': chapter_arabic or '',
                        'arabic_text': row_data.get(col_map.get('arabic', ''), ''),
                        'urdu_text': row_data.get(col_map.get('urdu', ''), ''),
                        'english_text': row_data.get(col_map.get('english', ''), ''),
                        'reference': row_data.get(col_map.get('reference', ''), ''),
                        'status': row_data.get(col_map.get('status', ''), ''),
                        'detailed_explanation': row_data.get(col_map.get('explanation', ''), ''),
                        'updated_by': request.user,
                    }

                    hadith_obj, created = Hadith.objects.update_or_create(
                        book=book,
                        hadith_number=hadith_no,
                        defaults=defaults,
                    )
                    if created:
                        created_count += 1
                    else:
                        updated_count += 1

            chapter = Chapter.objects.get(book=book, chapter_number=chapter_number)
            changed = False
            if chapter_english and chapter.chapter_english != chapter_english:
                chapter.chapter_english = chapter_english
                changed = True
            if chapter_urdu and chapter.chapter_urdu != chapter_urdu:
                chapter.chapter_urdu = chapter_urdu
                changed = True
            if chapter_arabic and chapter.chapter_arabic != chapter_arabic:
                chapter.chapter_arabic = chapter_arabic
                changed = True
            if changed:
                chapter.save()

            obj.result_log = (
                f"Successfully processed.\n"
                f"Book: {book.name}\n"
                f"Chapter: {chapter_number}. {chapter_english or chapter.chapter_english}\n"
                f"Hadiths: {len(rows)} processed "
                f"({updated_count} updated, {created_count} new)"
            )

        except Exception as e:
            obj.result_log = f"Error: {str(e)}"

        ModelClass.objects.filter(pk=obj.pk).update(result_log=obj.result_log)

